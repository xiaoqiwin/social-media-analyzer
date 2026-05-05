# -*- coding: utf-8 -*-
"""
舆情分析报告生成模块
功能：基于数据库统计结果，自动生成舆情分析报告
作者：Python自动化报告导师
日期：2026-05-02
安装命令：pip install python-docx
"""

import pymysql
import time
import os
from datetime import datetime, timedelta
from typing import Dict, List, Tuple, Any, Optional
import re

# 导入项目模块
from config import get_db_connection
import utils

# 尝试导入python-docx
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ 未安装python-docx，请运行: pip install python-docx")

# 获取日志记录器
logger = utils.get_logger(__name__)


def ensure_reports_dir() -> str:
    """
    确保报告输出目录存在

    Returns:
        str: 报告目录路径
    """
    return utils.ensure_output_dir(['reports'])


def get_overview(conn) -> Dict[str, Any]:
    """
    获取数据概览

    Args:
        conn: 数据库连接对象

    Returns:
        Dict[str, Any]: 概览数据
    """
    cursor = None
    try:
        cursor = conn.cursor()

        # 1. 从hot_events统计
        sql_events = """
        SELECT 
            COUNT(*) as total_events,
            COUNT(CASE WHEN DATE(crawl_date) = CURDATE() THEN 1 END) as today_events
        FROM hot_events
        """
        cursor.execute(sql_events)
        events_data = cursor.fetchone()

        # 2. 从comments统计
        sql_comments = """
        SELECT 
            COUNT(*) as total_comments,
            COUNT(CASE WHEN sentiment_label IS NOT NULL THEN 1 END) as analyzed_comments
        FROM comments
        """
        cursor.execute(sql_comments)
        comments_data = cursor.fetchone()

        # 3. 从sentiment_results统计
        sql_reports = """
        SELECT COUNT(*) as total_reports FROM sentiment_results
        """
        cursor.execute(sql_reports)
        reports_data = cursor.fetchone()

        return {
            'events': {
                'total': events_data['total_events'] or 0,
                'today': events_data['today_events'] or 0
            },
            'comments': {
                'total': comments_data['total_comments'] or 0,
                'analyzed': comments_data['analyzed_comments'] or 0
            },
            'reports': {
                'total': reports_data['total_reports'] or 0
            },
            'analyzed_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }

    except Exception as e:
        logger.error(f"获取数据概览失败: {e}")
        return {
            'events': {'total': 0, 'today': 0},
            'comments': {'total': 0, 'analyzed': 0},
            'reports': {'total': 0},
            'analyzed_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        }
    finally:
        if cursor:
            cursor.close()


def get_top5(conn) -> List[Dict[str, Any]]:
    """
    获取热点TOP5榜单

    Args:
        conn: 数据库连接对象

    Returns:
        List[Dict[str, Any]]: TOP5热点列表
    """
    cursor = None
    try:
        cursor = conn.cursor()

        sql = """
        SELECT title, hot_value, source, crawl_date
        FROM hot_events
        WHERE DATE(crawl_date) = CURDATE()
        ORDER BY hot_value DESC
        LIMIT 5
        """
        cursor.execute(sql)
        results = cursor.fetchall()

        if not results:
            # 如果没有今天的数据，获取所有数据
            sql = """
            SELECT title, hot_value, source, crawl_date
            FROM hot_events
            ORDER BY hot_value DESC
            LIMIT 5
            """
            cursor.execute(sql)
            results = cursor.fetchall()

        top5 = []
        for i, row in enumerate(results, 1):
            top5.append({
                'rank': i,
                'title': row['title'],
                'hot_value': row['hot_value'],
                'source': row['source'],
                'date': str(row['crawl_date'])
            })

        return top5

    except Exception as e:
        logger.error(f"获取TOP5榜单失败: {e}")
        return []
    finally:
        if cursor:
            cursor.close()


def get_sentiment_insight(conn) -> Dict[str, Any]:
    """
    获取情感洞察

    Args:
        conn: 数据库连接对象

    Returns:
        Dict[str, Any]: 情感洞察数据
    """
    cursor = None
    try:
        cursor = conn.cursor()

        # 1. 整体情感分布
        sql_overall = """
        SELECT 
            SUM(positive_count) as total_positive,
            SUM(negative_count) as total_negative,
            SUM(neutral_count) as total_neutral,
            SUM(total_count) as total_comments
        FROM sentiment_results
        WHERE DATE(analyzed_at) = CURDATE()
        """
        cursor.execute(sql_overall)
        overall = cursor.fetchone()

        if not overall or overall['total_comments'] == 0:
            # 如果没有今天的数据，使用所有数据
            sql_overall = """
            SELECT 
                SUM(positive_count) as total_positive,
                SUM(negative_count) as total_negative,
                SUM(neutral_count) as total_neutral,
                SUM(total_count) as total_comments
            FROM sentiment_results
            """
            cursor.execute(sql_overall)
            overall = cursor.fetchone()

        # 计算百分比
        total = overall['total_comments'] or 1
        positive_pct = (overall['total_positive'] or 0) / total * 100
        negative_pct = (overall['total_negative'] or 0) / total * 100
        neutral_pct = (overall['total_neutral'] or 0) / total * 100

        # 2. 最正面事件
        sql_most_positive = """
        SELECT 
            e.id,
            e.title,
            s.positive_count,
            s.negative_count,
            s.neutral_count,
            s.total_count,
            (s.positive_count * 1.0 / NULLIF(s.total_count, 0)) as positive_ratio
        FROM hot_events e
        JOIN sentiment_results s ON e.id = s.event_id
        WHERE DATE(s.analyzed_at) = CURDATE() AND s.total_count > 0
        ORDER BY positive_ratio DESC
        LIMIT 1
        """
        cursor.execute(sql_most_positive)
        most_positive = cursor.fetchone()

        # 3. 最负面事件
        sql_most_negative = """
        SELECT 
            e.id,
            e.title,
            s.positive_count,
            s.negative_count,
            s.neutral_count,
            s.total_count,
            (s.negative_count * 1.0 / NULLIF(s.total_count, 0)) as negative_ratio
        FROM hot_events e
        JOIN sentiment_results s ON e.id = s.event_id
        WHERE DATE(s.analyzed_at) = CURDATE() AND s.total_count > 0
        ORDER BY negative_ratio DESC
        LIMIT 1
        """
        cursor.execute(sql_most_negative)
        most_negative = cursor.fetchone()

        return {
            'overall': {
                'positive': positive_pct,
                'negative': negative_pct,
                'neutral': neutral_pct,
                'total': total
            },
            'most_positive': {
                'title': most_positive['title'] if most_positive else None,
                'positive_ratio': most_positive['positive_ratio'] * 100 if most_positive else 0,
                'positive_count': most_positive['positive_count'] if most_positive else 0,
                'total_count': most_positive['total_count'] if most_positive else 0
            } if most_positive else None,
            'most_negative': {
                'title': most_negative['title'] if most_negative else None,
                'negative_ratio': most_negative['negative_ratio'] * 100 if most_negative else 0,
                'negative_count': most_negative['negative_count'] if most_negative else 0,
                'total_count': most_negative['total_count'] if most_negative else 0
            } if most_negative else None
        }

    except Exception as e:
        logger.error(f"获取情感洞察失败: {e}")
        return {
            'overall': {'positive': 0, 'negative': 0, 'neutral': 0, 'total': 0},
            'most_positive': None,
            'most_negative': None
        }
    finally:
        if cursor:
            cursor.close()


def get_trend_insight(conn) -> Dict[str, Any]:
    """
    获取时间趋势总结

    Args:
        conn: 数据库连接对象

    Returns:
        Dict[str, Any]: 时间趋势数据
    """
    cursor = None
    try:
        cursor = conn.cursor()

        # 获取今日事件中情感波动最大的事件
        sql_fluctuation = """
        SELECT 
            e.id,
            e.title,
            e.hot_value,
            s.positive_count,
            s.negative_count,
            s.neutral_count,
            s.total_count,
            ABS(s.positive_count - s.negative_count) as sentiment_diff
        FROM hot_events e
        JOIN sentiment_results s ON e.id = s.event_id
        WHERE DATE(s.analyzed_at) = CURDATE() 
            AND DATE(e.crawl_date) = CURDATE()
            AND s.total_count > 10
        ORDER BY sentiment_diff DESC, e.hot_value DESC
        LIMIT 1
        """
        cursor.execute(sql_fluctuation)
        most_balanced = cursor.fetchone()

        sql_fluctuation_min = """
        SELECT 
            e.id,
            e.title,
            e.hot_value,
            s.positive_count,
            s.negative_count,
            s.neutral_count,
            s.total_count,
            ABS(s.positive_count - s.negative_count) as sentiment_diff
        FROM hot_events e
        JOIN sentiment_results s ON e.id = s.event_id
        WHERE DATE(s.analyzed_at) = CURDATE() 
            AND DATE(e.crawl_date) = CURDATE()
            AND s.total_count > 10
        ORDER BY sentiment_diff ASC, e.hot_value DESC
        LIMIT 1
        """
        cursor.execute(sql_fluctuation_min)
        most_polarized = cursor.fetchone()

        return {
            'most_balanced': {
                'title': most_balanced['title'] if most_balanced else None,
                'positive': most_balanced['positive_count'] if most_balanced else 0,
                'negative': most_balanced['negative_count'] if most_balanced else 0,
                'diff': most_balanced['sentiment_diff'] if most_balanced else 0,
                'total': most_balanced['total_count'] if most_balanced else 0
            } if most_balanced else None,
            'most_polarized': {
                'title': most_polarized['title'] if most_polarized else None,
                'positive': most_polarized['positive_count'] if most_polarized else 0,
                'negative': most_polarized['negative_count'] if most_polarized else 0,
                'diff': most_polarized['sentiment_diff'] if most_polarized else 0,
                'total': most_polarized['total_count'] if most_polarized else 0
            } if most_polarized else None
        }

    except Exception as e:
        logger.error(f"获取时间趋势总结失败: {e}")
        return {
            'most_balanced': None,
            'most_polarized': None
        }
    finally:
        if cursor:
            cursor.close()


def get_opinion_suggestions(conn) -> List[Dict[str, Any]]:
    """
    获取舆情建议

    Args:
        conn: 数据库连接对象

    Returns:
        List[Dict[str, Any]]: 舆情建议列表
    """
    cursor = None
    suggestions = []

    try:
        cursor = conn.cursor()

        # 获取今日有情感分析的事件
        sql = """
        SELECT 
            e.id,
            e.title,
            e.hot_value,
            s.positive_count,
            s.negative_count,
            s.neutral_count,
            s.total_count
        FROM hot_events e
        JOIN sentiment_results s ON e.id = s.event_id
        WHERE DATE(s.analyzed_at) = CURDATE() 
            AND s.total_count > 5
        ORDER BY e.hot_value DESC
        """
        cursor.execute(sql)
        events = cursor.fetchall()

        for event in events:
            title = event['title']
            hot_value = event['hot_value']
            positive = event['positive_count'] or 0
            negative = event['negative_count'] or 0
            neutral = event['neutral_count'] or 0
            total = event['total_count'] or 0

            # 规则1: 负面情绪占主导
            if total > 0 and negative > positive * 2:
                suggestions.append({
                    'event': title[:50] + "..." if len(title) > 50 else title,
                    'type': 'negative_dominant',
                    'description': f'该事件负面评论({negative}条)是正面评论({positive}条)的{negative / positive:.1f}倍，负面情绪占主导，建议关注',
                    'hot_value': hot_value
                })

            # 规则2: 高热度正面关注
            if hot_value > 1000000 and positive > negative:
                suggestions.append({
                    'event': title[:50] + "..." if len(title) > 50 else title,
                    'type': 'high_hot_positive',
                    'description': f'该事件热度高达{hot_value:,}，且正面评论({positive}条)多于负面评论({negative}条)，获得广泛正面关注',
                    'hot_value': hot_value
                })

            # 规则3: 争议性事件（正负面相当）
            if total > 20 and 0.8 < positive / negative < 1.2 and positive > 5 and negative > 5:
                suggestions.append({
                    'event': title[:50] + "..." if len(title) > 50 else title,
                    'type': 'controversial',
                    'description': f'该事件正面评论({positive}条)和负面评论({negative}条)数量相当，争议性较大',
                    'hot_value': hot_value
                })

            # 规则4: 中性评论过多
            if total > 10 and neutral / total > 0.7:
                suggestions.append({
                    'event': title[:50] + "..." if len(title) > 50 else title,
                    'type': 'neutral_heavy',
                    'description': f'该事件中性评论占比过高({neutral / total * 100:.1f}%)，用户情绪表达不明确',
                    'hot_value': hot_value
                })

        return suggestions

    except Exception as e:
        logger.error(f"获取舆情建议失败: {e}")
        return []
    finally:
        if cursor:
            cursor.close()


def generate_md(data: Dict[str, Any]) -> str:
    """
    生成Markdown格式报告

    Args:
        data: 报告数据

    Returns:
        str: Markdown格式报告
    """
    # 报告模板
    template = """# 社交媒体舆情分析报告

**生成时间**: {analyzed_at}  
**数据范围**: 微博热点话题

---

## 1. 数据概览

| 指标类别 | 总数 | 今日新增 |
|---------|------|----------|
| 热点事件 | {total_events} | {today_events} |
| 用户评论 | {total_comments} | 已分析 {analyzed_comments} |
| 情感报告 | {total_reports} | 份 |

> 数据说明：本报告基于{crawl_date}采集的数据生成。

---

## 2. 热点TOP5榜单

{top5_list}

---

## 3. 情感洞察

### 3.1 整体情感分布

- **正面评论**: {positive_pct:.1f}% ({positive_count:,}条)
- **负面评论**: {negative_pct:.1f}% ({negative_count:,}条)
- **中性评论**: {neutral_pct:.1f}% ({neutral_count:,}条)
- **总计**: {total_sentiment:,}条评论

### 3.2 最正面事件

{most_positive_info}

### 3.3 最负面事件

{most_negative_info}

---

## 4. 时间趋势总结

### 4.1 情感最平衡事件

{most_balanced_info}

### 4.2 情感最对立事件

{most_polarized_info}

---

## 5. 舆情建议

{suggestions_list}

---

## 6. 分析图表

已生成以下可视化图表，请查看对应文件：

{charts_list}

---

*报告生成系统: 社交媒体热点话题分析系统*  
*生成时间: {generated_time}*
"""

    # 准备数据
    overview = data.get('overview', {})
    top5 = data.get('top5', [])
    sentiment = data.get('sentiment', {})
    trend = data.get('trend', {})
    suggestions = data.get('suggestions', [])

    # 今日日期
    today = datetime.now().strftime('%Y-%m-%d')

    # 生成TOP5列表
    top5_list = ""
    for i, item in enumerate(top5, 1):
        top5_list += f"{i}. **{item['title']}**  \n"
        top5_list += f"   热度: {item['hot_value']:,} | 来源: {item['source']} | 日期: {item['date']}\n\n"

    if not top5_list:
        top5_list = "暂无热点事件数据"

    # 情感洞察
    sentiment_overall = sentiment.get('overall', {})
    most_positive = sentiment.get('most_positive')
    most_negative = sentiment.get('most_negative')

    most_positive_info = ""
    if most_positive and most_positive.get('title'):
        most_positive_info = f"**{most_positive['title']}**  \n"
        most_positive_info += f"正面评论占比: {most_positive['positive_ratio']:.1f}% "
        most_positive_info += f"({most_positive['positive_count']}/{most_positive['total_count']}条)"
    else:
        most_positive_info = "暂无正面占主导的事件"

    most_negative_info = ""
    if most_negative and most_negative.get('title'):
        most_negative_info = f"**{most_negative['title']}**  \n"
        most_negative_info += f"负面评论占比: {most_negative['negative_ratio']:.1f}% "
        most_negative_info += f"({most_negative['negative_count']}/{most_negative['total_count']}条)"
    else:
        most_negative_info = "暂无负面占主导的事件"

    # 时间趋势
    most_balanced = trend.get('most_balanced')
    most_polarized = trend.get('most_polarized')

    most_balanced_info = ""
    if most_balanced and most_balanced.get('title'):
        most_balanced_info = f"**{most_balanced['title']}**  \n"
        most_balanced_info += f"正面: {most_balanced['positive']}条, "
        most_balanced_info += f"负面: {most_balanced['negative']}条, "
        most_balanced_info += f"差值: {most_balanced['diff']}条"
    else:
        most_balanced_info = "暂无数据"

    most_polarized_info = ""
    if most_polarized and most_polarized.get('title'):
        most_polarized_info = f"**{most_polarized['title']}**  \n"
        most_polarized_info += f"正面: {most_polarized['positive']}条, "
        most_polarized_info += f"负面: {most_polarized['negative']}条, "
        most_polarized_info += f"差值: {most_polarized['diff']}条"
    else:
        most_polarized_info = "暂无数据"

    # 舆情建议
    suggestions_list = ""
    if suggestions:
        for i, suggestion in enumerate(suggestions, 1):
            suggestions_list += f"{i}. **{suggestion['event']}**  \n"
            suggestions_list += f"   {suggestion['description']}  \n\n"
    else:
        suggestions_list = "暂无特殊舆情建议"

    # 图表列表
    charts_dir = "output/charts/"
    chart_files = [
        "wordcloud.html",  # 词云图
        "hot_trend.html",  # 热度趋势图
        "sentiment_pie.html",  # 情感分布饼图
        "top10_bar.html",  # TOP10柱状图
        "combined.html"  # 组合图表
    ]

    charts_list = ""
    for chart_file in chart_files:
        chart_path = os.path.join(charts_dir, chart_file)
        if os.path.exists(chart_path):
            charts_list += f"- ✅ {chart_file} (已生成)\n"
        else:
            charts_list += f"- ⏳ {chart_file} (未生成)\n"

    # 格式化模板
    md_content = template.format(
        analyzed_at=overview.get('analyzed_at', today + ' 00:00:00'),
        total_events=overview.get('events', {}).get('total', 0),
        today_events=overview.get('events', {}).get('today', 0),
        total_comments=overview.get('comments', {}).get('total', 0),
        analyzed_comments=overview.get('comments', {}).get('analyzed', 0),
        total_reports=overview.get('reports', {}).get('total', 0),
        crawl_date=today,
        top5_list=top5_list,
        positive_pct=sentiment_overall.get('positive', 0),
        negative_pct=sentiment_overall.get('negative', 0),
        neutral_pct=sentiment_overall.get('neutral', 0),
        positive_count=int(sentiment_overall.get('total', 0) * sentiment_overall.get('positive', 0) / 100),
        negative_count=int(sentiment_overall.get('total', 0) * sentiment_overall.get('negative', 0) / 100),
        neutral_count=int(sentiment_overall.get('total', 0) * sentiment_overall.get('neutral', 0) / 100),
        total_sentiment=sentiment_overall.get('total', 0),
        most_positive_info=most_positive_info,
        most_negative_info=most_negative_info,
        most_balanced_info=most_balanced_info,
        most_polarized_info=most_polarized_info,
        suggestions_list=suggestions_list,
        charts_list=charts_list,
        generated_time=datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    )

    return md_content


def save_md_report(md_content: str) -> str:
    """
    保存Markdown报告

    Args:
        md_content: Markdown内容

    Returns:
        str: 保存的文件路径
    """
    try:
        output_dir = ensure_reports_dir()
        file_path = os.path.join(output_dir, "report.md")

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

        logger.info(f"Markdown报告已保存: {file_path}")
        return file_path

    except Exception as e:
        logger.error(f"保存Markdown报告失败: {e}")
        return ""


def export_docx_report(md_content: str) -> str:
    """
    导出Word文档报告

    Args:
        md_content: Markdown内容

    Returns:
        str: 保存的Word文件路径
    """
    if not DOCX_AVAILABLE:
        logger.error("未安装python-docx，无法生成Word文档")
        return ""

    try:
        output_dir = ensure_reports_dir()
        docx_path = os.path.join(output_dir, "report.docx")

        # 创建文档
        doc = Document()

        # 设置中文字体
        try:
            # 设置正文样式
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            style.font.size = Pt(10.5)

            # 设置标题样式
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = '黑体'
            title_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            title_style.font.size = Pt(16)
            title_style.font.bold = True

            # 设置一级标题样式
            heading1_style = doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            heading1_style.font.name = '黑体'
            heading1_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            heading1_style.font.size = Pt(14)
            heading1_style.font.bold = True

            # 设置二级标题样式
            heading2_style = doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            heading2_style.font.name = '黑体'
            heading2_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            heading2_style.font.size = Pt(12)
            heading2_style.font.bold = True
        except Exception as e:
            logger.warning(f"设置中文字体失败: {e}")

        # 解析Markdown内容
        lines = md_content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                # 空行
                doc.add_paragraph()
                continue

            if line.startswith('# '):
                # 主标题
                title_text = line[2:].strip()
                p = doc.add_paragraph(title_text, style='CustomTitle')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif line.startswith('## '):
                # 一级标题
                heading_text = line[3:].strip()
                p = doc.add_paragraph(heading_text, style='CustomHeading1')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            elif line.startswith('### '):
                # 二级标题
                heading_text = line[4:].strip()
                p = doc.add_paragraph(heading_text, style='CustomHeading2')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            elif line.startswith('---'):
                # 水平线
                doc.add_paragraph("─" * 50)

            elif line.startswith('> '):
                # 引用
                quote_text = line[2:].strip()
                p = doc.add_paragraph(quote_text)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)

            elif line.startswith('- '):
                # 列表项
                list_text = line[2:].strip()
                p = doc.add_paragraph(list_text, style='List Bullet')

            elif line.startswith('|'):
                # 表格行
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells and '---' not in line:
                    # 如果是表格内容行
                    row_text = ' | '.join(cells)
                    p = doc.add_paragraph(row_text)

            elif line.startswith('1. '):
                # 有序列表
                list_text = line[3:].strip()
                p = doc.add_paragraph(list_text, style='List Number')

            elif '**' in line:
                # 粗体文本
                # 简单处理粗体标记
                text = line.replace('**', '')
                p = doc.add_paragraph(text)

            else:
                # 普通段落
                p = doc.add_paragraph(line)

        # 添加页脚
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | 社交媒体热点话题分析系统"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存文档
        doc.save(docx_path)
        logger.info(f"Word文档报告已保存: {docx_path}")

        return docx_path

    except Exception as e:
        logger.error(f"生成Word文档失败: {e}")
        return ""


def print_report_summary(data: Dict[str, Any]) -> None:
    """
    在控制台打印报告摘要

    Args:
        data: 报告数据
    """
    print("\n" + "=" * 60)
    print("📊 舆情分析报告摘要")
    print("=" * 60)

    # 数据概览
    overview = data.get('overview', {})
    print("📈 数据概览:")
    print(f"  热点事件: {overview.get('events', {}).get('total', 0)} 个")
    print(f"  今日新增: {overview.get('events', {}).get('today', 0)} 个")
    print(f"  用户评论: {overview.get('comments', {}).get('total', 0)} 条")
    print(f"  已分析评论: {overview.get('comments', {}).get('analyzed', 0)} 条")

    # TOP5热点
    top5 = data.get('top5', [])
    print(f"\n🔥 热点TOP5:")
    for item in top5:
        print(f"  {item['rank']}. {item['title'][:30]}... (热度: {item['hot_value']:,})")

    # 情感分布
    sentiment = data.get('sentiment', {})
    sentiment_overall = sentiment.get('overall', {})
    print(f"\n😊 情感分布:")
    print(f"  正面: {sentiment_overall.get('positive', 0):.1f}%")
    print(f"  负面: {sentiment_overall.get('negative', 0):.1f}%")
    print(f"  中性: {sentiment_overall.get('neutral', 0):.1f}%")

    # 舆情建议
    suggestions = data.get('suggestions', [])
    print(f"\n💡 舆情建议: {len(suggestions)} 条")
    for i, suggestion in enumerate(suggestions[:3], 1):  # 只显示前3条
        print(f"  {i}. {suggestion['description'][:50]}...")

    if len(suggestions) > 3:
        print(f"  ... 还有 {len(suggestions) - 3} 条建议")

    print("=" * 60)


def run_reporter() -> bool:
    """
    运行报告生成模块的主函数

    这个函数将被main.py调用

    Returns:
        bool: 是否运行成功
    """
    if not DOCX_AVAILABLE:
        print("❌ 未安装python-docx，Word文档功能不可用")
        print("请运行以下命令安装:")
        print("  pip install python-docx")
        # 继续运行，只是不生成Word文档

    print("📄 启动舆情分析报告生成模块")
    print("=" * 60)

    conn = None
    try:
        # 获取数据库连接
        conn = get_db_connection()

        # 收集数据
        print("1. 正在收集数据概览...")
        overview = get_overview(conn)

        print("2. 正在获取热点TOP5...")
        top5 = get_top5(conn)

        print("3. 正在分析情感洞察...")
        sentiment = get_sentiment_insight(conn)

        print("4. 正在总结时间趋势...")
        trend = get_trend_insight(conn)

        print("5. 正在生成舆情建议...")
        suggestions = get_opinion_suggestions(conn)

        # 整合数据
        data = {
            'overview': overview,
            'top5': top5,
            'sentiment': sentiment,
            'trend': trend,
            'suggestions': suggestions
        }

        # 在控制台打印摘要
        print_report_summary(data)

        # 生成Markdown报告
        print("\n6. 正在生成Markdown报告...")
        md_content = generate_md(data)
        md_path = save_md_report(md_content)

        if md_path:
            print(f"✅ Markdown报告已保存: {md_path}")
        else:
            print("❌ Markdown报告保存失败")

        # 生成Word文档报告
        if DOCX_AVAILABLE:
            print("\n7. 正在生成Word文档报告...")
            docx_path = export_docx_report(md_content)

            if docx_path:
                print(f"✅ Word文档报告已保存: {docx_path}")
            else:
                print("❌ Word文档报告生成失败")
        else:
            print("\n⚠️ 跳过Word文档生成（python-docx未安装）")

        # 输出总结
        print("\n" + "=" * 60)
        print("🎯 报告生成完成")
        print("=" * 60)
        print(f"📁 输出目录: {os.path.abspath('output/reports')}")
        print(f"📄 生成文件:")
        print(f"  - report.md (Markdown格式)")
        if DOCX_AVAILABLE:
            print(f"  - report.docx (Word格式)")
        print("=" * 60)

        return True

    except Exception as e:
        logger.error(f"报告生成模块运行失败: {e}")
        print(f"❌ 错误: {e}")
        return False
    finally:
        if conn:
            conn.close()


if __name__ == "__main__":
    """
    直接运行此模块时，执行报告生成
    """
    print("📄 舆情分析报告生成模块独立运行")
    print("=" * 60)

    success = run_reporter()
    if success:
        print("✅ 报告生成模块执行完成！")
    else:
        print("❌ 报告生成模块执行失败！")